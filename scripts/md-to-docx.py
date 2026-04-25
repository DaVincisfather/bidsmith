#!/usr/bin/env python
"""Convert markdown to docx for E2E testing of the markitdown parser.

Supports: H1-H4, bullets, numbered lists, tables (with header bolding),
bold (**), italic (* and _), paragraphs.

Usage: python md-to-docx.py <input.md> <output.docx>
"""
import re
import sys
from pathlib import Path

from docx import Document


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)")


def add_runs(paragraph, text: str) -> None:
    """Split text into python-docx runs honoring **bold**, *italic*, _italic_."""
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") or token.startswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, start_idx):
    """Parse a markdown table starting at lines[start_idx]. Returns (rows, end_idx)."""
    rows = []
    idx = start_idx
    sep_pattern = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        line = lines[idx].strip()
        if sep_pattern.match(line):
            idx += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        idx += 1
    return rows, idx


def render_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    bullet_re = re.compile(r"^\s*[-*]\s+(.*)$")
    number_re = re.compile(r"^\s*\d+\.\s+(.*)$")
    sep_pattern = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")
    block_start = re.compile(r"^(#{1,6}\s|\s*[-*]\s|\s*\d+\.\s|\|)")

    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        if not stripped.strip():
            i += 1
            continue

        m = heading_re.match(stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        if stripped.lstrip().startswith("|") and i + 1 < n and sep_pattern.match(lines[i + 1]):
            rows, i = parse_table(lines, i)
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        text = row[c_idx] if c_idx < len(row) else ""
                        add_runs(cell.paragraphs[0], text)
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue

        m = bullet_re.match(stripped)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, m.group(1))
            i += 1
            continue

        m = number_re.match(stripped)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_runs(para, m.group(1))
            i += 1
            continue

        para_lines = [stripped.strip()]
        i += 1
        while i < n and lines[i].strip() and not block_start.match(lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        para = doc.add_paragraph()
        add_runs(para, " ".join(para_lines))

    doc.save(str(out_path))


def main() -> None:
    if len(sys.argv) != 3:
        print("Usage: python md-to-docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)
    in_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    if not in_path.exists():
        print(f"Input file not found: {in_path}", file=sys.stderr)
        sys.exit(1)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    md_text = in_path.read_text(encoding="utf-8")
    render_docx(md_text, out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
